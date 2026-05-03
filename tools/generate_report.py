# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT = "output/doc/report_tmm26_variant25.docx"


def font(run, size=12, name="Times New Roman", bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def para(doc, text="", align=None, bold=False, first=True):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    if first:
        p.paragraph_format.first_line_indent = Cm(1.25)
    if align:
        p.alignment = align
    r = p.add_run(text)
    font(r, bold=bold)
    return p


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    for r in p.runs:
        font(r, 14 if level == 1 else 13, bold=True)
    return p


def code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    font(r, 9, "Courier New")
    return p


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text)
    font(r, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell_text(t.rows[0].cells[i], h, True)
        shade(t.rows[0].cells[i], "D9EAF7")
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cell_text(cells[i], value)
    return t


def page(doc):
    doc.add_page_break()


def add_title(doc):
    lines = [
        "Министерство науки и высшего образования Российской Федерации",
        "Федеральное государственное автономное образовательное учреждение высшего образования",
        "«Московский государственный технический университет имени Н.Э. Баумана",
        "(национальный исследовательский университет)»",
        "(МГТУ им. Н.Э. Баумана)",
        "",
        "Факультет «Информатика и системы управления»",
        "Кафедра «Системы обработки информации и управления»",
        "",
        "",
    ]
    for line in lines:
        para(doc, line, WD_ALIGN_PARAGRAPH.CENTER, first=False)
    para(doc, "Отчет по ДЗ", WD_ALIGN_PARAGRAPH.CENTER, True, False)
    para(doc, "«Обнаружение и отслеживание наличия мобильника в руках человека»", WD_ALIGN_PARAGRAPH.CENTER, True, False)
    para(doc, "по дисциплине «Технология мультимедиа»", WD_ALIGN_PARAGRAPH.CENTER, False, False)
    para(doc, "", first=False)
    para(doc, "Вариант: 25", WD_ALIGN_PARAGRAPH.CENTER, True, False)
    para(doc, "", first=False)
    para(doc, "", first=False)
    para(doc, "Выполнил:", first=False)
    para(doc, "студент(ка) группы № ____________    ____________________________", first=False)
    para(doc, "подпись, дата", first=False)
    para(doc, "", first=False)
    para(doc, "Проверил:", first=False)
    para(doc, "к.т.н., доц., Г.И. Афанасьев", first=False)
    para(doc, "подпись, дата", first=False)
    for _ in range(8):
        para(doc, "", first=False)
    para(doc, "2026 г.", WD_ALIGN_PARAGRAPH.CENTER, False, False)
    page(doc)


def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(12)

    add_title(doc)

    heading(doc, "Содержание")
    for item in [
        "1. Титульный лист",
        "2. Задание",
        "3. Введение",
        "4. Конструкторская часть",
        "4.1. Прототипы и аналоги",
        "4.2. Функциональная модульная схема программы",
        "4.3. Схема общего алгоритма работы программы",
        "4.4. Описание алгоритмов ключевых модулей",
        "4.5. Ключевые скрины экрана смартфона",
        "4.6. Нюансы реализации и мобильная нейронная сеть",
        "4.7. Системные требования и инструкция по запуску",
        "5. Исследовательская часть",
        "6. Заключение",
        "7. Библиографический список",
        "Приложение А. Основные файлы проекта",
    ]:
        para(doc, item, first=False)
    page(doc)

    heading(doc, "2. Задание")
    for text in [
        "Вариант 25 домашнего задания ТММ 26: обнаружение и отслеживание наличия мобильника в руках человека в видеопотоке с видеокамеры в реальном масштабе времени.",
        "Цель работы - разработать мобильное Android-приложение, которое получает видеопоток с камеры смартфона, выполняет локальное распознавание объектов и визуализирует факт наличия мобильного телефона в зоне рук человека.",
        "Программа должна запускаться из Android Studio на подключенном смартфоне, работать без удаленного сервера, показывать видеопоток, рамки обнаруженных объектов и итоговое состояние анализа.",
    ]:
        para(doc, text)

    heading(doc, "3. Введение")
    for text in [
        "Задача обнаружения телефона в руках относится к классу задач компьютерного зрения для анализа поведения человека в видеопотоке. Она актуальна для систем контроля безопасности, анализа внимания водителя, мониторинга рабочих мест и учебных лабораторных стендов.",
        "Классический подход включает получение кадров видеопотока, детекцию объектов, пространственный анализ взаимного положения объектов, временное сглаживание результата и визуализацию состояния. Для мобильных устройств особенно важны малый размер модели, высокая скорость инференса и автономная работа.",
        "В данной работе выбран локальный вариант решения на базе Android CameraX и TensorFlow Lite. CameraX отвечает за получение кадров с камеры, а модель EfficientDet-Lite0 выполняет детекцию классов COCO, включая person и cell phone.",
        "Наличие телефона в руках определяется не только фактом обнаружения телефона, но и дополнительной логикой связи рамки телефона с рамкой человека и зоной предполагаемого положения рук. Такой подход делает решение простым, воспроизводимым и пригодным для защиты на физическом смартфоне.",
    ]:
        para(doc, text)
    page(doc)

    heading(doc, "4. Конструкторская часть")
    heading(doc, "4.1. Прототипы и аналоги", 2)
    table(doc, ["Источник", "Назначение", "Использование в проекте"], [
        ["TensorFlow Lite Android object detection example", "Локальная детекция объектов на Android", "Идея TFLite-инференса на кадрах камеры"],
        ["Android CameraX documentation", "Preview и ImageAnalysis", "Получение кадров в реальном времени"],
        ["EfficientDet-Lite model family", "Легкие модели object detection", "Модель EfficientDet-Lite0 с метаданными"],
        ["COCO dataset classes", "Набор классов object detection", "Классы person и cell phone"],
    ])
    para(doc, "Близкие аналоги обычно решают общую задачу object detection. В разработанном проекте добавлен специализированный слой PhoneInHandAnalyzer, который принимает решение о наличии телефона в руках по геометрии рамок и истории кадров.")

    heading(doc, "4.2. Функциональная модульная схема программы", 2)
    para(doc, "Функциональная схема приложения:")
    code(doc, "Камера смартфона -> CameraX Preview/ImageAnalysis -> Bitmap кадра -> TensorFlow Lite ObjectDetector -> список рамок -> PhoneInHandAnalyzer -> OverlayView и статус на экране")
    table(doc, ["Модуль", "Файл", "Функция"], [
        ["Главный экран", "MainActivity.kt", "Инициализация камеры, модели, анализатора и UI"],
        ["Детектор объектов", "TensorFlow Lite Task Vision", "Поиск person и cell phone на кадре"],
        ["Анализ состояния", "PhoneInHandAnalyzer.kt", "Связь телефона с зоной рук и сглаживание"],
        ["Визуализация", "OverlayView.kt", "Рисование рамок поверх камеры"],
        ["Модель данных", "DetectionModels.kt", "Структуры DetectionBox, PhoneState, PhoneAnalysis"],
    ])
    page(doc)

    heading(doc, "4.3. Схема общего алгоритма работы программы", 2)
    for step in [
        "Запросить разрешение CAMERA у пользователя.",
        "Запустить CameraX Preview для отображения видеопотока.",
        "Запустить ImageAnalysis с политикой STRATEGY_KEEP_ONLY_LATEST.",
        "Скопировать RGBA-кадр в Bitmap и повернуть изображение по ориентации камеры.",
        "Передать кадр в TensorFlow Lite ObjectDetector.",
        "Выбрать лучшие рамки классов cell phone и person.",
        "Оценить пересечение и положение телефона относительно рамки человека.",
        "Сгладить результат по нескольким кадрам.",
        "Вывести рамки, confidence, время инференса и итоговый статус на экран.",
    ]:
        para(doc, step)
    code(doc, 'while camera_is_active:\n    frame = get_latest_frame()\n    detections = object_detector.detect(frame)\n    phone = best_detection(detections, "cell phone")\n    person = best_detection(detections, "person")\n    state = analyze_spatial_relation(phone, person)\n    state = smooth_state_by_frame_history(state)\n    draw_overlay(frame, phone, person, state)')

    heading(doc, "4.4. Описание алгоритмов ключевых модулей", 2)
    para(doc, "MainActivity создает интерфейс, запускает камеру и передает кадры в модель. Для снижения задержки используется однопоточный executor и стратегия пропуска старых кадров.")
    code(doc, "val imageAnalyzer = ImageAnalysis.Builder()\n    .setBackpressureStrategy(ImageAnalysis.STRATEGY_KEEP_ONLY_LATEST)\n    .setOutputImageFormat(ImageAnalysis.OUTPUT_IMAGE_FORMAT_RGBA_8888)\n    .build()\nimageAnalyzer.setAnalyzer(cameraExecutor) { image -> detect(image) }")
    para(doc, "PhoneInHandAnalyzer выбирает самый уверенный телефон и наиболее связанного с ним человека. Затем вычисляются нормированные координаты центра телефона в рамке человека, пересечение рамок и принадлежность телефона к зоне вероятного положения рук.")
    code(doc, "val normalizedX = (phone.centerX() - person.left) / person.width()\nval normalizedY = (phone.centerY() - person.top) / person.height()\nval plausibleArmHeight = normalizedY in 0.28f..0.92f\nval nearBodySide = normalizedX in -0.12f..0.35f || normalizedX in 0.65f..1.12f\nreturn insideExpandedPerson && plausibleArmHeight && nearBodySide")
    para(doc, "Временное сглаживание снижает мигание результата. PHONE_IN_HAND включается после двух последовательных положительных кадров, а PHONE_VISIBLE удерживается до нескольких отрицательных кадров.")
    code(doc, "if (rawInHand) {\n    positiveStreak++\n    negativeStreak = 0\n} else {\n    negativeStreak++\n    positiveStreak = 0\n}\nval state = when {\n    positiveStreak >= 2 -> PHONE_IN_HAND\n    negativeStreak >= 4 -> PHONE_VISIBLE\n    else -> lastState\n}")
    page(doc)

    heading(doc, "4.5. Ключевые скрины экрана смартфона", 2)
    para(doc, "После запуска на смартфоне необходимо вставить реальные скриншоты работы приложения. Предусмотрены следующие визуальные состояния:")
    table(doc, ["Скриншот", "Состояние", "Описание"], [
        ["Скрин 1", "Телефон не обнаружен", "Камера работает, статус зеленый, рамка телефона отсутствует"],
        ["Скрин 2", "Телефон найден", "Красная рамка вокруг телефона, связь с руками не подтверждена"],
        ["Скрин 3", "Телефон в руке", "Телефон расположен в зоне рук человека, статус красный"],
        ["Скрин 4", "Несколько кадров подряд", "Статус не мигает благодаря временному сглаживанию"],
    ])
    for label in [
        "Место для вставки скриншота 1: экран без телефона.",
        "Место для вставки скриншота 2: телефон обнаружен в кадре.",
        "Место для вставки скриншота 3: телефон в руке человека.",
    ]:
        para(doc, label, first=False)
        for _ in range(4):
            para(doc, "", first=False)
    page(doc)

    heading(doc, "4.6. Нюансы реализации и описание мобильной нейронной сети", 2)
    for text in [
        "В проекте используется модель EfficientDet-Lite0 в формате TensorFlow Lite. Это легкая модель object detection, оптимизированная для мобильных и embedded-устройств. Файл модели размещен в app/src/main/assets/efficientdet-lite0.tflite и содержит метаданные классов.",
        "EfficientDet использует эффективную архитектуру с масштабированием backbone, feature pyramid и prediction heads. Версия Lite0 является самой компактной в семействе и подходит для real-time сценариев на смартфоне.",
        "Модель используется без дополнительного обучения, поскольку класс cell phone уже присутствует в COCO-разметке. Поверх результата нейронной сети добавлен геометрический модуль, так как сама модель не возвращает отношение “объект находится в руке”.",
        "Видеопоток обрабатывается локально. Кадры не отправляются в сеть, что уменьшает задержку, повышает приватность и упрощает демонстрацию программы на преподавательском компьютере.",
    ]:
        para(doc, text)

    heading(doc, "4.7. Системные требования и инструкция по запуску", 2)
    table(doc, ["Компонент", "Требование"], [
        ["ОС разработки", "Windows 10/11"],
        ["IDE", "Android Studio"],
        ["JDK", "JBR/JDK 17"],
        ["Android SDK", "compileSdk 35, minSdk 24"],
        ["Смартфон", "Android 7.0 или выше, камера, USB debugging"],
        ["Библиотеки", "CameraX, TensorFlow Lite Task Vision"],
    ])
    for step in [
        "Открыть папку проекта Phone-In-Hand-Detector в Android Studio.",
        "Дождаться Gradle Sync.",
        "Подключить смартфон по USB и включить режим разработчика.",
        "Выбрать конфигурацию app и нажать Run.",
        "Разрешить приложению доступ к камере.",
        "Навести камеру на человека с телефоном в руке и проверить изменение статуса.",
    ]:
        para(doc, step)
    page(doc)

    heading(doc, "5. Исследовательская часть")
    para(doc, "Для проверки работоспособности программы предлагается провести эксперимент в четырех условиях: телефон отсутствует, телефон лежит отдельно от человека, телефон находится в руке, телефон частично перекрыт пальцами или корпусом. Для каждого условия фиксируется не менее 30 секунд видеопотока.")
    table(doc, ["Условие", "Ожидаемый результат", "Метрика"], [
        ["Телефона нет", "NO_PHONE", "Доля кадров без ложного PHONE_IN_HAND"],
        ["Телефон на столе", "PHONE_VISIBLE", "Доля кадров без ложной связи с руками"],
        ["Телефон в правой руке", "PHONE_IN_HAND", "Доля кадров с верным обнаружением"],
        ["Телефон в левой руке", "PHONE_IN_HAND", "Доля кадров с верным обнаружением"],
        ["Частичное перекрытие", "PHONE_IN_HAND или PHONE_VISIBLE", "Устойчивость при снижении confidence"],
    ])
    para(doc, "Пример таблицы экспериментальных результатов после демонстрации:")
    table(doc, ["Сценарий", "Кадров всего", "Верных состояний", "Точность, %", "Среднее время инференса, мс"], [
        ["Телефона нет", "300", "___", "___", "___"],
        ["Телефон на столе", "300", "___", "___", "___"],
        ["Телефон в руке", "300", "___", "___", "___"],
        ["Перекрытие телефона", "300", "___", "___", "___"],
    ])
    para(doc, "На практике точность зависит от освещения, размера телефона на кадре, расстояния до человека и положения руки. При хорошем освещении и видимом телефоне модель стабильно выделяет класс cell phone. При сильном перекрытии confidence снижается, поэтому состояние может перейти из PHONE_IN_HAND в PHONE_VISIBLE.")
    page(doc)

    heading(doc, "6. Заключение")
    for text in [
        "В ходе работы разработан Android-проект для обнаружения и отслеживания наличия мобильного телефона в руках человека в видеопотоке с камеры. Программа использует CameraX для получения кадров и TensorFlow Lite Task Vision для локальной детекции объектов.",
        "Реализованы модули визуализации рамок, анализа пространственной связи телефона с человеком и временного сглаживания результата. На экране смартфона отображаются видеопоток, рамки обнаруженных объектов, confidence и итоговый статус.",
        "Достоинством решения является работа без сервера и возможность демонстрации непосредственно на смартфоне. Ограничение связано с тем, что отношение “телефон в руке” определяется эвристически по рамкам объектов. Для повышения точности можно добавить отдельную модель keypoint detection кистей или дообучить модель на специализированном наборе данных.",
    ]:
        para(doc, text)

    heading(doc, "7. Библиографический список")
    for ref in [
        "1. Android Developers. CameraX overview. https://developer.android.com/media/camera/camerax",
        "2. TensorFlow Lite. Object detection with TensorFlow Lite. https://www.tensorflow.org/lite/examples/object_detection/overview",
        "3. TensorFlow Lite Task Library. Vision APIs. https://www.tensorflow.org/lite/inference_with_metadata/task_library/overview",
        "4. Tan M., Pang R., Le Q.V. EfficientDet: Scalable and Efficient Object Detection. https://arxiv.org/abs/1911.09070",
        "5. COCO Dataset. Common Objects in Context. https://cocodataset.org/",
        "6. Android Developers. Configure your build. https://developer.android.com/build",
    ]:
        para(doc, ref, first=False)
    page(doc)

    heading(doc, "Приложение А. Основные файлы проекта")
    table(doc, ["Файл", "Назначение"], [
        ["settings.gradle", "Описание модулей проекта и репозиториев Gradle"],
        ["build.gradle", "Версии Android Gradle Plugin и Kotlin"],
        ["app/build.gradle", "Android-конфигурация, зависимости CameraX и TensorFlow Lite"],
        ["MainActivity.kt", "Камера, анализ кадров и UI"],
        ["PhoneInHandAnalyzer.kt", "Алгоритм принятия решения"],
        ["OverlayView.kt", "Отрисовка рамок поверх камеры"],
        ["efficientdet-lite0.tflite", "Мобильная нейронная сеть object detection"],
    ])
    para(doc, "Полный код проекта предоставляется отдельно в папке Android Studio. В отчет включены только ключевые фрагменты, отражающие алгоритм работы программы.")

    for p in doc.paragraphs:
        is_code = any(r.font.name == "Courier New" for r in p.runs)
        p.paragraph_format.line_spacing = 1.0 if is_code else 1.5
        p.paragraph_format.space_after = Pt(0)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
